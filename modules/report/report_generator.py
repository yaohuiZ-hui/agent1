"""
审计报告生成模块 - 报告生成器

在学员完成全部实操任务后，自动汇总全流程数据，生成
符合银行业标准的《数据库安全运维与加固报告》。

报告内容：
1. 概述（学员信息、任务时间线）
2. 漏洞列表（SQL注入点、弱口令、过度授权）
3. 修复措施（代码修复、WAF配置、权限变更）
4. 权限变更清单
5. 基线检查结果
6. 数据恢复时间点
7. 综合评分与等级
"""
import os
import json
from typing import Optional, Dict, Any, List
from datetime import datetime
from dataclasses import dataclass, field, asdict

from config.settings import get_config, Config
from core.database_connector import DatabaseConnector
from core.agent_orchestrator import STORY_BRANCHES


@dataclass
class ReportSection:
    """报告章节"""
    title: str
    content: List[str]
    severity: str = "info"


class ReportGenerator:
    """
    报告生成器

    收集全流程数据，生成符合银行业标准的PDF/HTML安全运维报告。
    """

    def __init__(self, db: DatabaseConnector, config: Optional[Config] = None):
        self._db = db
        self._config = config or get_config()

    # ──────────────────────────────────────────
    # 数据收集
    # ──────────────────────────────────────────

    def collect_all_data(self, student_id: int = 1, terminal_state: dict = None) -> Dict[str, Any]:
        """收集全流程操作数据"""
        # 学员信息
        student_rows = self._db.execute_sqlite(
            "SELECT * FROM student_state WHERE id = ?", (student_id,)
        )
        student = dict(student_rows[0]) if student_rows else {}
        student["completed_tasks"] = json.loads(student.get("completed_tasks", "[]")) if student else []

        # 任务记录
        tasks = self._db.execute_sqlite(
            "SELECT * FROM task_records WHERE student_id = ? ORDER BY started_at", (student_id,)
        )
        task_list = [dict(t) for t in tasks]

        # 命令历史
        commands = self._db.execute_sqlite(
            "SELECT * FROM command_history WHERE student_id = ? ORDER BY executed_at", (student_id,)
        )
        cmd_list = [dict(c) for c in commands]

        # 漏洞记录
        vulns = self._db.execute_sqlite(
            "SELECT * FROM vulnerability_records WHERE student_id = ? ORDER BY found_at", (student_id,)
        )
        vuln_list = [dict(v) for v in vulns]

        # 权限变更
        perm_changes = self._db.execute_sqlite(
            "SELECT * FROM permission_changes WHERE student_id = ? ORDER BY changed_at", (student_id,)
        )
        perm_list = [dict(p) for p in perm_changes]

        # 基线检查结果
        baseline = self._db.execute_sqlite(
            "SELECT * FROM baseline_results WHERE student_id = ? ORDER BY checked_at", (student_id,)
        )
        baseline_list = [dict(b) for b in baseline]

        # 审计日志
        audit = self._db.execute_sqlite(
            "SELECT * FROM audit_logs WHERE student_id = ? ORDER BY logged_at", (student_id,)
        )
        audit_list = [dict(a) for a in audit]

        return {
            "student": student,
            "tasks": task_list,
            "commands": cmd_list,
            "vulnerabilities": vuln_list,
            "permission_changes": perm_list,
            "baseline_results": baseline_list,
            "audit_logs": audit_list,
            "terminal_state": terminal_state or {},
            "generated_at": datetime.now().isoformat(),
        }

    # ──────────────────────────────────────────
    # 报告章节生成
    # ──────────────────────────────────────────

    def _build_overview_section(self, data: Dict[str, Any]) -> ReportSection:
        """生成概述章节"""
        student = data.get("student", {})
        tasks = data.get("tasks", [])

        completed_count = sum(1 for t in tasks if t.get("status") == "completed")
        failed_count = sum(1 for t in tasks if t.get("status") == "failed")

        content = [
            f"学员名称: {student.get('student_name', 'N/A')}",
            f"报告生成时间: {data.get('generated_at', datetime.now().isoformat())[:19]}",
            f"任务完成情况: {completed_count}个完成, {failed_count}个失败",
            f"完成度: {completed_count / max(len(tasks), 1) * 100:.0f}%",
            f"故事分支: {student.get('current_branch', 'N/A')}",
            f"最终阶段: {student.get('story_phase', 'N/A')}",
        ]

        return ReportSection("1. 概述", content, "info")

    def _get_vulnerability_list(self, data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """从所有可用数据源确定完整漏洞列表和修复状态"""
        # 获取已完成任务
        tasks = data.get("tasks", [])
        completed_task_ids = {t.get("task_id") for t in tasks if t.get("status") == "completed"}

        # 获取终端权限状态（用于精确判断）
        state = data.get("terminal_state", {})
        perms = state.get("permissions", {})
        if not perms:
            try:
                raw = self._db.load_state("terminal_state")
                if raw:
                    saved_state = json.loads(raw)
                    perms = saved_state.get("permissions", {})
            except Exception:
                pass

        # 构建完整的已知漏洞清单（覆盖本场景全部安全风险）
        # 各项修复状态根据学员实际完成的任务动态判定
        # 基线修复任务ID集合
        baseline_fix_done = bool({"fix_root_remote", "fix_anonymous", "fix_testuser_revoke", "fix_devuser_revoke", "fix_appuser_revoke"} & completed_task_ids)
        all_revoke_done = "fix_testuser_revoke" in completed_task_ids and "fix_devuser_revoke" in completed_task_ids and "fix_appuser_revoke" in completed_task_ids
        all_known_vulns = [
            {
                "vuln_type": "匿名用户",
                "endpoint": "''@localhost",
                "severity": "high",
                "is_fixed": 1 if "fix_anonymous" in completed_task_ids else 0,
                "fixed_method": "需执行 DROP USER ''@'localhost' 删除匿名用户" if "fix_anonymous" not in completed_task_ids else "已执行 DROP USER ''@'localhost'",
            },
            {
                "vuln_type": "root远程登录",
                "endpoint": "root@%",
                "severity": "critical",
                "is_fixed": 1 if "fix_root_remote" in completed_task_ids else 0,
                "fixed_method": "需禁止root远程登录（DELETE FROM mysql.user WHERE user='root' AND host='%'）" if "fix_root_remote" not in completed_task_ids else "已禁止root远程登录",
            },
            {
                "vuln_type": "弱口令",
                "endpoint": "test_user, dev_user等",
                "severity": "high",
                "is_fixed": 1 if baseline_fix_done else 0,
                "fixed_method": "需使用 ALTER USER 修改弱口令账号密码" if not baseline_fix_done else "已修复弱口令账号密码",
            },
            {
                "vuln_type": "测试账号过度授权",
                "endpoint": "test_user@%",
                "severity": "critical",
                "is_fixed": 1 if "fix_testuser_revoke" in completed_task_ids else 0,
                "fixed_method": "需 REVOKE DELETE,DROP,FILE ON *.* FROM 'test_user'@'%'" if "fix_testuser_revoke" not in completed_task_ids else "已撤销测试账号过度权限",
            },
            {
                "vuln_type": "开发账号过度授权",
                "endpoint": "dev_user@%",
                "severity": "high",
                "is_fixed": 1 if "fix_devuser_revoke" in completed_task_ids else 0,
                "fixed_method": "需 REVOKE DELETE,DROP,ALTER ON *.* FROM 'dev_user'@'%'" if "fix_devuser_revoke" not in completed_task_ids else "已撤销开发账号过度权限",
            },
            {
                "vuln_type": "应用账号过度授权",
                "endpoint": "app_user@192.168.%",
                "severity": "high",
                "is_fixed": 1 if "fix_appuser_revoke" in completed_task_ids else 0,
                "fixed_method": "需 REVOKE DELETE ON *.* FROM 'app_user'@'192.168.%'" if "fix_appuser_revoke" not in completed_task_ids else "已撤销应用账号过度权限",
            },
            {
                "vuln_type": "FILE权限滥用",
                "endpoint": "test_user, dev_user",
                "severity": "high",
                "is_fixed": 1 if all_revoke_done else 0,
                "fixed_method": "需 REVOKE FILE ON *.* FROM 相关用户" if not all_revoke_done else "已撤销FILE权限",
            },
            {
                "vuln_type": "Union注入",
                "endpoint": "/api/user/profile",
                "severity": "critical",
                "is_fixed": 1 if "fix_vulnerable_code" in completed_task_ids else 0,
                "fixed_method": "需将字符串拼接SQL改为参数化查询" if "fix_vulnerable_code" not in completed_task_ids else "已使用参数化查询修复代码",
            },
            {
                "vuln_type": "盲注",
                "endpoint": "/api/search/product",
                "severity": "high",
                "is_fixed": 1 if "configure_waf" in completed_task_ids else 0,
                "fixed_method": "需部署WAF拦截 + 参数化查询" if "configure_waf" not in completed_task_ids else "已配置WAF规则拦截",
            },
            {
                "vuln_type": "SSL/TLS传输加密",
                "endpoint": "数据库连接",
                "severity": "info",
                "is_fixed": 1,
                "fixed_method": "✓ SSL已启用 (have_ssl=YES)，无需修复",
            },
            {
                "vuln_type": "审计日志未启用",
                "endpoint": "general_log",
                "severity": "medium",
                "is_fixed": 1 if "configure_waf" in completed_task_ids or "fix_vulnerable_code" in completed_task_ids else 0,
                "fixed_method": "需 SET GLOBAL general_log=ON 启用审计日志" if "configure_waf" not in completed_task_ids and "fix_vulnerable_code" not in completed_task_ids else "建议启用审计日志",
            },
            {
                "vuln_type": "默认端口(3306)",
                "endpoint": "MySQL端口",
                "severity": "info",
                "is_fixed": 1,
                "fixed_method": "提示：默认端口3306需在my.cnf中修改并重启MySQL服务（非紧急项）",
            },
            {
                "vuln_type": "密码过期策略未设置",
                "endpoint": "default_password_lifetime",
                "severity": "medium",
                "is_fixed": 1 if baseline_fix_done else 0,
                "fixed_method": "需 SET GLOBAL default_password_lifetime=180" if not baseline_fix_done else "密码策略已加固",
            },
        ]

        return all_known_vulns

    def _build_vulnerability_section(self, data: Dict[str, Any]) -> ReportSection:
        """生成漏洞列表章节"""
        vulns = self._get_vulnerability_list(data)

        content = [f"共发现 {len(vulns)} 个安全漏洞"]
        for v in vulns:
            status = "✓ 已修复" if v.get("is_fixed") else "✗ 未修复"
            content.append(f"  [{v.get('severity', 'medium').upper()}] {v.get('endpoint', 'N/A')} - {v.get('vuln_type', 'N/A')} [{status}]")
            if v.get("fixed_method"):
                content.append(f"    修复方式: {v['fixed_method']}")

        return ReportSection("2. 漏洞列表及修复措施", content,
                             "critical" if any(not v.get("is_fixed") for v in vulns) else "info")

    def _build_permission_section(self, data: Dict[str, Any]) -> ReportSection:
        """生成权限变更章节（根据修复任务完成状态动态更新）"""
        changes = data.get("permission_changes", [])
        state = data.get("terminal_state", {})

        # 获取已完成任务
        tasks = data.get("tasks", [])
        completed_task_ids = {t.get("task_id") for t in tasks if t.get("status") == "completed"}
        perm_fixed = "fix_testuser_revoke" in completed_task_ids or "fix_devuser_revoke" in completed_task_ids or "fix_appuser_revoke" in completed_task_ids

        if not changes and state:
            perms = state.get("permissions", {})
            db_changes = state.get("db_changes", [])
            if db_changes:
                changes = db_changes
                # 同步更新状态：若权限修复任务已完成，待处理项改为已处理
                if perm_fixed:
                    for c in changes:
                        if c.get("status") in ("pending", "待处理"):
                            c["status"] = "applied"
            else:
                changes = []
                for key, info in perms.items():
                    u, h = (key.split("@") + ["%"])[:2] if "@" in key else [key, "%"]
                    gp = info.get("global_privs", []); role = info.get("role", "")
                    status = "applied" if perm_fixed else "pending"
                    if ("DELETE" in gp or "DROP" in gp) and role in ("测试账号","开发账号","应用账号"):
                        changes.append({"change_type": "REVOKE", "target_user": u, "target_host": h, "privilege": "DELETE,DROP", "status": status})
                    if "FILE" in gp and role not in ("管理员",):
                        changes.append({"change_type": "REVOKE", "target_user": u, "target_host": h, "privilege": "FILE", "status": status})
                    if u == "root" and h == "%":
                        changes.append({"change_type": "DELETE", "target_user": "root(远程)", "target_host": "%", "privilege": "禁止远程登录", "status": status})
                    if u == "''":
                        changes.append({"change_type": "DROP", "target_user": "''(匿名)", "target_host": "localhost", "privilege": "匿名用户", "status": status})
                if not changes:
                    changes.append({"change_type": "-", "target_user": "-", "target_host": "-", "privilege": "无待处理项", "status": "applied"})

        content = [f"权限变更记录: {len(changes)} 项"]
        for c in changes:
            status_icon = "✓" if c.get('status') == 'applied' else "○" if c.get('status') == 'pending' else "✗"
            content.append(f"  [{status_icon}] {c.get('change_type', 'N/A')} "
                           f"用户 {c.get('target_user', 'N/A')}@{c.get('target_host', 'N/A')} - {c.get('privilege', 'N/A')}")

        return ReportSection("3. 权限变更清单", content, "info")

    def _build_baseline_section(self, data: Dict[str, Any]) -> ReportSection:
        """生成基线检查结果章节"""
        from modules.toolkit.baseline_checker import BaselineChecker
        checker = BaselineChecker()
        baseline_result = checker.run_all_checks()

        summary = baseline_result["summary"]
        content = [
            f"CIS MySQL Benchmark 基线检查",
            f"总检查项: {summary['total']}",
            f"通过: {summary['passed']}",
            f"失败: {summary['failed']}",
            f"安全评分: {summary['score']}/100 (等级: {summary['grade']})",
            "",
            "详细结果:",
        ]
        for detail in baseline_result["details"]:
            icon = "✓" if detail["status"] == "pass" else "✗"
            content.append(f"  [{icon}] [{detail['severity'].upper()}] {detail['check_id']} {detail['check_name']}")
            if detail["status"] == "fail":
                content.append(f"    实际值: {detail['actual_value']}")
                content.append(f"    期望值: {detail['expected_value']}")

        return ReportSection("4. 基线检查结果", content,
                             "critical" if summary["score"] < 60 else "warning" if summary["score"] < 80 else "info")

    def _build_recovery_section(self, data: Dict[str, Any]) -> ReportSection:
        """生成数据恢复章节（根据修复任务完成状态动态更新）"""
        tasks = data.get("tasks", [])
        completed_task_ids = {t.get("task_id") for t in tasks if t.get("status") == "completed"}

        # 判断数据恢复任务完成情况
        recovery_done = all(tid in completed_task_ids for tid in ["restore_full_backup", "apply_binlog_pitr", "verify_data_integrity"])
        recovery_partial = any(tid in completed_task_ids for tid in ["restore_full_backup", "apply_binlog_pitr", "verify_data_integrity"])
        branch = data.get("student", {}).get("current_branch", "")

        if recovery_done:
            content = [
                "恢复类型: 时间点恢复 (PITR - Point-in-Time Recovery)",
                "恢复工具: Percona XtraBackup + mysqlbinlog",
                "",
                "恢复过程:",
                "  1. ✓ 使用 XtraBackup 准备全量备份 (2024-01-15_full.xb)",
                "  2. ✓ 恢复数据文件到 /var/lib/mysql/",
                "  3. ✓ 使用 mysqlbinlog 回放二进制日志到指定时间点",
                "  4. ✓ 校验数据完整性 (checksum验证通过)",
                "",
                "恢复结果:",
                "  ✓ 成功恢复 1500 条交易记录",
                "  ✓ 数据完整性校验通过",
                "  ⚠ 发现 3 条 account_no 字段异常数据 (怀疑SQL注入篡改)",
            ]
            severity = "info"
        elif recovery_partial:
            content = [
                "恢复类型: 时间点恢复 (PITR - Point-in-Time Recovery)",
                "恢复状态: 部分完成（以下为已完成步骤）",
                "",
                "恢复过程:",
            ]
            if "restore_full_backup" in completed_task_ids:
                content.append("  1. ✓ 使用 XtraBackup 准备全量备份")
                content.append("  2. ✓ 恢复数据文件到 /var/lib/mysql/")
            else:
                content.append("  1. ✗ 全量备份恢复 — 未执行")
            if "apply_binlog_pitr" in completed_task_ids:
                content.append("  3. ✓ 使用 mysqlbinlog 回放二进制日志到指定时间点")
            else:
                content.append("  3. ✗ Binlog时间点恢复 — 未执行")
            if "verify_data_integrity" in completed_task_ids:
                content.append("  4. ✓ 校验数据完整性")
            else:
                content.append("  4. ✗ 数据完整性校验 — 未执行")
            content.append("")
            content.append("⚠ 数据恢复尚未全部完成，请继续执行剩余恢复步骤")
            severity = "warning"
        else:
            content = [
                "恢复类型: 时间点恢复 (PITR - Point-in-Time Recovery)",
                "恢复状态: 未执行",
                "",
                "说明: 学员当前处于「{}」分支，未选择数据恢复路线或未执行恢复操作。".format(branch),
                "如需数据恢复，请依次执行:",
                "  1. xtrabackup --prepare --apply-log-only --target-dir=/backup/full/",
                "  2. xtrabackup --copy-back --target-dir=/backup/full/ --datadir=/var/lib/mysql/",
                "  3. mysqlbinlog --stop-datetime='恢复时间点' /backup/binlog/mysql-bin.000012 | mysql -u root -p",
                "  4. CHECKSUM TABLE core_bank.trade_flow 校验数据完整性",
            ]
            severity = "warning"

        return ReportSection("5. 数据恢复情况", content, severity)

    def _build_recommendation_section(self) -> ReportSection:
        """生成修复建议章节"""
        content = [
            "1. 立即执行:",
            "   • 所有生产环境使用参数化查询，杜绝字符串拼接SQL",
            "   • 撤销所有非管理员的全局DELETE/DROP/FILE权限",
            "   • 禁止root远程登录，仅允许localhost访问",
            "   • 删除匿名用户，启用密码强度验证插件",
            "",
            "2. 短期整改（1-2周）:",
            "   • 部署WAF (ModSecurity + OWASP CRS规则集)",
            "   • 启用审计日志 (general_log) 和慢查询日志",
            "   • 修改默认MySQL端口(3306)",
            "   • 启用SSL/TLS加密传输",
            "   • 设置密码过期策略 (180天)",
            "",
            "3. 长期规划:",
            "   • 部署数据库审计系统 (McAfee/MariaDB Audit Plugin)",
            "   • 建立定期安全基线检查机制 (每月)",
            "   • 实施数据库权限季度审查制度",
            "   • 建立灾备恢复演练机制 (每季度)",
        ]
        return ReportSection("6. 修复加固建议", content, "info")

    def _build_score_section(self, data: Dict[str, Any]) -> ReportSection:
        """生成综合评分章节"""
        from core.agent_orchestrator import get_orchestrator
        orchestrator = get_orchestrator(self._db)
        score_data = orchestrator.calculate_final_score()
        student = data.get("student", {})

        content = [
            f"学员: {student.get('student_name', 'N/A')}",
            f"基础分: {score_data['base_score']}/100",
            f"失败扣分: -{score_data['fail_penalty']}",
            f"重试扣分: -{score_data['retry_penalty']}",
            f"漏洞修复加分: +{score_data['vuln_bonus']}",
            f"最终得分: {score_data['final_score']}/100",
            f"等级评定: {score_data['grade']}",
        ]

        severity = "info"
        if score_data["final_score"] < 60:
            severity = "critical"
        elif score_data["final_score"] < 80:
            severity = "warning"

        return ReportSection("7. 综合评分与等级", content, severity)

    def _get_branch_failure_counts(self, data: Dict[str, Any]) -> Dict[str, int]:
        """从学员状态中解析各分支错误次数（ADR-0001 D10）。"""
        student = data.get("student", {})
        raw = student.get("branch_failed_counts", "")
        try:
            counts = json.loads(raw) if isinstance(raw, str) and raw else (raw or {})
        except Exception:
            counts = {}
        return counts if isinstance(counts, dict) else {}

    def _build_branch_failure_section(self, data: Dict[str, Any]) -> ReportSection:
        """生成各分支错误次数章节（text 报告使用）。"""
        counts = self._get_branch_failure_counts(data)
        content = ["各分支错误次数统计（操作失误累计）:", ""]
        if not counts:
            content.append("  (无任何分支产生操作失误)")
        else:
            for branch in ("baseline", "recovery", "sqli"):
                cnt = counts.get(branch, 0)
                maxf = STORY_BRANCHES.get(branch, {}).get("max_failures", "-")
                title = STORY_BRANCHES.get(branch, {}).get("title", branch)
                content.append(f"  [{branch}] {title}: {cnt} 次 / 预算 {maxf} 次")
            for branch, cnt in counts.items():
                if branch not in ("baseline", "recovery", "sqli"):
                    content.append(f"  [{branch}]: {cnt} 次")
        return ReportSection("8. 各分支错误次数", content,
                             "warning" if counts else "info")

    # ──────────────────────────────────────────
    # 报告生成
    # ──────────────────────────────────────────

    def generate_text_report(self, student_id: int = 1, terminal_state: dict = None) -> str:
        """生成文本格式报告"""
        data = self.collect_all_data(student_id, terminal_state)
        sections = [
            self._build_overview_section(data),
            self._build_vulnerability_section(data),
            self._build_permission_section(data),
            self._build_baseline_section(data),
            self._build_recovery_section(data),
            self._build_recommendation_section(),
            self._build_score_section(data),
            self._build_branch_failure_section(data),
        ]

        lines = [
            "=" * 70,
            "  数据库安全运维与加固报告",
            f"  生成日期: {datetime.now().strftime('%Y年%m月%d日')}",
            "=" * 70,
            "",
        ]

        for section in sections:
            lines.append(f"{'=' * 70}")
            lines.append(f"  {section.title}")
            lines.append(f"{'=' * 70}")
            for line in section.content:
                lines.append(f"  {line}")
            lines.append("")

        # 签名
        lines.append("-" * 70)
        lines.append("  本报告由 数据库安全运维智能体(Agent 1) 自动生成")
        lines.append("  符合银行业安全运维标准")
        lines.append("=" * 70)

        return "\n".join(lines)

    def generate_html_report(self, student_id: int = 1, terminal_state: dict = None) -> str:
        """生成HTML格式报告"""
        data = self.collect_all_data(student_id, terminal_state)
        score_data = self._get_score_data(student_id)

        # 构建HTML
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>数据库安全运维与加固报告</title>
<style>
    body {{ font-family: 'SimSun', 'Microsoft YaHei', sans-serif; margin: 40px; color: #333; }}
    h1 {{ text-align: center; color: #1a237e; font-size: 24px; margin-bottom: 5px; }}
    h2 {{ color: #1a237e; border-bottom: 2px solid #1a237e; padding-bottom: 5px; margin-top: 30px; }}
    h3 {{ color: #283593; }}
    .header {{ text-align: center; margin-bottom: 30px; }}
    .header p {{ color: #666; font-size: 14px; }}
    .section {{ margin: 20px 0; padding: 15px; background: #f8f9fa; border-radius: 5px; }}
    .critical {{ border-left: 4px solid #d32f2f; }}
    .warning {{ border-left: 4px solid #f57c00; }}
    .info {{ border-left: 4px solid #1976d2; }}
    .score-box {{ text-align: center; padding: 30px; background: linear-gradient(135deg, #1a237e, #283593); color: white; border-radius: 10px; margin: 20px 0; }}
    .score-box .score {{ font-size: 48px; font-weight: bold; }}
    .score-box .grade {{ font-size: 24px; margin-top: 10px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
    th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
    th {{ background: #1a237e; color: white; }}
    tr:nth-child(even) {{ background: #f2f2f2; }}
    .vuln-critical {{ color: #d32f2f; font-weight: bold; }}
    .vuln-high {{ color: #f57c00; font-weight: bold; }}
    .vuln-medium {{ color: #fbc02d; }}
    .status-fixed {{ color: #388e3c; }}
    .status-unfixed {{ color: #d32f2f; }}
    .recommendation {{ background: #e8f5e9; padding: 10px; margin: 5px 0; border-radius: 3px; }}
</style>
</head>
<body>
<div class="header">
    <h1>数据库安全运维与加固报告</h1>
    <p>报告编号: SEC-{datetime.now().strftime('%Y%m%d')}-{student_id:03d}</p>
    <p>生成日期: {datetime.now().strftime('%Y年%m月%d日')}</p>
    <p>密级: 内部秘密</p>
</div>
"""
        # 综合评分
        html += f"""
<div class="score-box">
    <div>综合评分</div>
    <div class="score">{score_data['final_score']}</div>
    <div class="grade">等级: {score_data['grade']}</div>
</div>
"""
        # 1. 概述
        student = data.get("student", {})
        html += f"""
<h2>1. 概述</h2>
<div class="section info">
    <table>
        <tr><td>学员名称</td><td>{student.get('student_name', 'N/A')}</td></tr>
        <tr><td>报告生成时间</td><td>{data.get('generated_at', '')[:19]}</td></tr>
        <tr><td>故事分支</td><td>{student.get('current_branch', 'N/A')}</td></tr>
        <tr><td>最终阶段</td><td>{student.get('story_phase', 'N/A')}</td></tr>
        <tr><td>完成任务</td><td>{len(student.get('completed_tasks', [])) if isinstance(student.get('completed_tasks'), list) else len(json.loads(student.get('completed_tasks', '[]')))} 个</td></tr>
        <tr><td>失败次数</td><td>{student.get('failed_count', 0)} 次</td></tr>
    </table>
</div>
"""
        # 2. 漏洞列表
        html += """
<h2>2. 漏洞列表及修复措施</h2>
<div class="section">
    <table>
        <tr><th>漏洞类型</th><th>端点/位置</th><th>严重级别</th><th>修复状态</th><th>修复方式</th></tr>
"""
        vulns = self._get_vulnerability_list(data)
        for v in vulns:
            sev_class = f"vuln-{v.get('severity', 'medium')}"
            status_text = "✓ 已修复" if v.get("is_fixed") else "✗ 未修复"
            status_class = "status-fixed" if v.get("is_fixed") else "status-unfixed"
            html += f"""
        <tr>
            <td>{v.get('vuln_type', 'N/A')}</td>
            <td>{v.get('endpoint', 'N/A')}</td>
            <td class="{sev_class}">{v.get('severity', 'medium').upper()}</td>
            <td class="{status_class}">{status_text}</td>
            <td>{v.get('fixed_method', '-')}</td>
        </tr>"""
        html += "</table></div>"

        # 3. 权限变更
        html += """
<h2>3. 权限变更清单</h2>
<div class="section">
    <table><tr><th>操作</th><th>用户</th><th>Host</th><th>权限/操作</th><th>状态</th></tr>
"""
        perm_changes = data.get("permission_changes", [])
        tasks = data.get("tasks", [])
        completed_task_ids = {t.get("task_id") for t in tasks if t.get("status") == "completed"}
        perm_fixed = "fix_testuser_revoke" in completed_task_ids or "fix_devuser_revoke" in completed_task_ids or "fix_appuser_revoke" in completed_task_ids
        if not perm_changes:
            status_text = "已处理" if perm_fixed else "待处理"
            perm_changes = [
                {"change_type": "REVOKE", "target_user": "test_user", "target_host": "%",
                 "privilege": "DELETE,DROP,FILE", "status": status_text},
                {"change_type": "DELETE", "target_user": "root(远程)", "target_host": "%",
                 "privilege": "禁止远程登录", "status": status_text},
                {"change_type": "REVOKE", "target_user": "dev_user", "target_host": "%",
                 "privilege": "DELETE,DROP,ALTER", "status": status_text},
                {"change_type": "REVOKE", "target_user": "app_user", "target_host": "192.168.%",
                 "privilege": "DELETE", "status": status_text},
                {"change_type": "DROP", "target_user": "匿名用户", "target_host": "localhost",
                 "privilege": "删除匿名用户", "status": status_text},
            ]
        for c in perm_changes:
            html += f"""<tr><td>{c.get('change_type','')}</td><td>{c.get('target_user','')}</td><td>{c.get('target_host','')}</td><td>{c.get('privilege','')}</td><td>{c.get('status','')}</td></tr>"""
        html += "</table></div>"


        # 5. 数据恢复
        html += """
<h2>4. 数据恢复情况</h2>
<div class="section info">
"""
        # 获取恢复任务完成状态
        recovery_tasks = data.get("tasks", [])
        recovery_completed = {t.get("task_id") for t in recovery_tasks if t.get("status") == "completed"}
        recovery_done = all(tid in recovery_completed for tid in ["restore_full_backup", "apply_binlog_pitr", "verify_data_integrity"])
        recovery_partial = any(tid in recovery_completed for tid in ["restore_full_backup", "apply_binlog_pitr", "verify_data_integrity"])
        branch_name = data.get("student", {}).get("current_branch", "未知")

        if recovery_done:
            html += """
    <p><strong>恢复类型:</strong> 时间点恢复 (PITR)</p>
    <p><strong>恢复工具:</strong> Percona XtraBackup + mysqlbinlog</p>
    <p><strong>恢复状态:</strong> ✓ 已完成</p>
    <h3>恢复过程</h3>
    <ol>
        <li>✓ 使用 XtraBackup 准备全量备份 (2024-01-15_full.xb)</li>
        <li>✓ 恢复数据文件到 /var/lib/mysql/</li>
        <li>✓ 使用 mysqlbinlog 回放二进制日志到指定时间点</li>
        <li>✓ 校验数据完整性 (checksum验证通过)</li>
    </ol>
    <h3>恢复结果</h3>
    <p>✓ 成功恢复 1500 条交易记录</p>
    <p>✓ 数据完整性校验通过</p>
    <p>⚠ 发现 3 条 account_no 字段异常数据</p>"""
        elif recovery_partial:
            html += """
    <p><strong>恢复类型:</strong> 时间点恢复 (PITR)</p>
    <p><strong>恢复状态:</strong> ⚠ 部分完成</p>
    <h3>恢复过程</h3>
    <ol>"""
            if "restore_full_backup" in recovery_completed:
                html += """
        <li>✓ 使用 XtraBackup 准备全量备份</li>
        <li>✓ 恢复数据文件到 /var/lib/mysql/</li>"""
            else:
                html += """
        <li>✗ 全量备份恢复 — 未执行</li>"""
            if "apply_binlog_pitr" in recovery_completed:
                html += """
        <li>✓ 使用 mysqlbinlog 回放二进制日志到指定时间点</li>"""
            else:
                html += """
        <li>✗ Binlog时间点恢复 — 未执行</li>"""
            if "verify_data_integrity" in recovery_completed:
                html += """
        <li>✓ 校验数据完整性</li>"""
            else:
                html += """
        <li>✗ 数据完整性校验 — 未执行</li>"""
            html += """
    </ol>
    <p>⚠ 数据恢复尚未全部完成，请继续执行剩余恢复步骤</p>"""
        else:
            html += f"""
    <p><strong>恢复类型:</strong> 时间点恢复 (PITR)</p>
    <p><strong>恢复状态:</strong> 未执行</p>
    <p>学员当前处于「{branch_name}」分支，未选择数据恢复路线或未执行恢复操作。</p>
    <p>如需数据恢复，请依次执行:</p>
    <ol>
        <li>xtrabackup --prepare --apply-log-only --target-dir=/backup/full/</li>
        <li>xtrabackup --copy-back --target-dir=/backup/full/ --datadir=/var/lib/mysql/</li>
        <li>mysqlbinlog --stop-datetime='恢复时间点' /backup/binlog/mysql-bin.000012 | mysql -u root -p</li>
        <li>CHECKSUM TABLE core_bank.trade_flow 校验数据完整性</li>
    </ol>"""

        html += "</div>"""

        # 6. 建议
        html += """
<h2>5. 修复加固建议</h2>
<div class="section">
    <div class="recommendation"><strong>1. 立即执行:</strong><br>
    • 所有生产环境使用参数化查询，杜绝字符串拼接SQL<br>
    • 撤销所有非管理员的全局DELETE/DROP/FILE权限<br>
    • 禁止root远程登录，仅允许localhost访问<br>
    • 删除匿名用户，启用密码强度验证插件</div>
    <div class="recommendation"><strong>2. 短期整改（1-2周）:</strong><br>
    • 部署WAF (ModSecurity + OWASP CRS规则集)<br>
    • 启用审计日志和慢查询日志<br>
    • 修改默认MySQL端口(3306)<br>
    • 启用SSL/TLS加密传输</div>
    <div class="recommendation"><strong>3. 长期规划:</strong><br>
    • 部署数据库审计系统<br>
    • 建立定期安全基线检查机制 (每月)<br>
    • 实施数据库权限季度审查制度</div>
</div>"""

        # 7. 综合评分
        html += f"""
<h2>6. 综合评分与等级</h2>
<div class="section">
    <table>
        <tr><td>基础分</td><td>{score_data['base_score']}/100</td></tr>
        <tr><td>失败扣分</td><td>-{score_data['fail_penalty']}</td></tr>
        <tr><td>重试扣分</td><td>-{score_data['retry_penalty']}</td></tr>
        <tr><td>漏洞修复加分</td><td>+{score_data['vuln_bonus']}</td></tr>
        <tr><td><strong>最终得分</strong></td><td><strong>{score_data['final_score']}/100</strong></td></tr>
        <tr><td><strong>等级评定</strong></td><td><strong>{score_data['grade']}</strong></td></tr>
    </table>
</div>"""
        # 7. 各分支错误次数（ADR-0001 D11）
        counts = self._get_branch_failure_counts(data)
        html += f"""
<h2>7. 各分支错误次数</h2>
<div class="section">
    <table>
        <tr><th>分支</th><th>名称</th><th>错误次数</th><th>错误预算</th></tr>"""
        if counts:
            for branch in ("baseline", "recovery", "sqli"):
                cnt = counts.get(branch, 0)
                maxf = STORY_BRANCHES.get(branch, {}).get("max_failures", "-")
                title = STORY_BRANCHES.get(branch, {}).get("title", branch)
                sev = "vuln-critical" if maxf != "-" and cnt >= maxf else ""
                html += f"""
        <tr><td>{branch}</td><td>{title}</td><td class="{sev}">{cnt} 次</td><td>{maxf} 次</td></tr>"""
            for branch, cnt in counts.items():
                if branch not in ("baseline", "recovery", "sqli"):
                    html += f"""
        <tr><td>{branch}</td><td>{branch}</td><td>{cnt} 次</td><td>-</td></tr>"""
        else:
            html += """
        <tr><td colspan="4">无任何分支产生操作失误</td></tr>"""
        html += """
    </table>
</div>"""

        html += """
<div style="text-align:center; margin-top:40px; color:#999; font-size:12px;">
    本报告由 数据库安全运维智能体(Agent 1) 自动生成<br>
    符合银行业安全运维标准
</div>
</body></html>"""

        return html

    def generate_docx_report(self, student_id: int = 1, terminal_state: dict = None) -> bytes:
        """生成Word(.docx)格式报告，样式与HTML报告保持一致"""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        import io

        data = self.collect_all_data(student_id, terminal_state)
        score_data = self._get_score_data(student_id)
        doc = Document()

        # ── 全局样式 ──
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.paragraph_format.space_after = Pt(4)

        for level in range(1, 4):
            hs = doc.styles[f'Heading {level}']
            hf = hs.font
            hf.name = 'Microsoft YaHei'
            hf.bold = True
            hf.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if level == 1:
                hf.size = Pt(18)
            elif level == 2:
                hf.size = Pt(14)
            else:
                hf.size = Pt(12)

        def set_cell_shading(cell, color_hex):
            """设置单元格底色"""
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        def set_cell_text(cell, text, bold=False, color=None, size=None):
            """设置单元格文字"""
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if bold:
                run.bold = True
            if color:
                run.font.color.rgb = color
            if size:
                run.font.size = Pt(size)

        # ══════════════════════════════════
        # 封面
        # ══════════════════════════════════
        for _ in range(6):
            doc.add_paragraph('')

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run('数据库安全运维与加固报告')
        title_run.font.size = Pt(28)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
        title_run.font.name = 'Microsoft YaHei'
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_paragraph('')
        info_lines = [
            f"报告编号: SEC-{datetime.now().strftime('%Y%m%d')}-{student_id:03d}",
            f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}",
            "密级: 内部秘密",
        ]
        for line in info_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            r.font.name = 'Microsoft YaHei'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_page_break()

        # ══════════════════════════════════
        # 综合评分 卡片
        # ══════════════════════════════════
        score_table = doc.add_table(rows=1, cols=1)
        score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        score_cell = score_table.cell(0, 0)
        set_cell_shading(score_cell, "1a237e")
        score_p = score_cell.paragraphs[0]
        score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        score_label_run = score_p.add_run('\n综合评分\n')
        score_label_run.font.size = Pt(14)
        score_label_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        score_label_run.font.name = 'Microsoft YaHei'
        score_label_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        score_val_run = score_p.add_run(f"{score_data['final_score']}\n")
        score_val_run.font.size = Pt(48)
        score_val_run.bold = True
        score_val_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        grade_run = score_p.add_run(f"等级: {score_data['grade']}\n\n")
        grade_run.font.size = Pt(20)
        grade_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        grade_run.font.name = 'Microsoft YaHei'
        grade_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_paragraph('')

        # ══════════════════════════════════
        # 1. 概述
        # ══════════════════════════════════
        doc.add_heading('1. 概述', level=1)
        student = data.get("student", {})
        overview_table = doc.add_table(rows=5, cols=2)
        overview_table.style = 'Table Grid'
        overview_rows = [
            ("学员名称", student.get('student_name', 'N/A')),
            ("报告生成时间", data.get('generated_at', '')[:19]),
            ("故事分支", student.get('current_branch', 'N/A')),
            ("最终阶段", student.get('story_phase', 'N/A')),
            ("已完成任务数", str(len(student.get('completed_tasks', [])) if isinstance(student.get('completed_tasks'), list) else len(json.loads(student.get('completed_tasks', '[]'))))),
        ]
        for i, (k, v) in enumerate(overview_rows):
            set_cell_text(overview_table.cell(i, 0), k, bold=True, size=11)
            set_cell_text(overview_table.cell(i, 1), v, size=11)
            set_cell_shading(overview_table.cell(i, 0), "e8eaf6")

        doc.add_paragraph('')

        # ══════════════════════════════════
        # 2. 漏洞列表及修复措施
        # ══════════════════════════════════
        doc.add_heading('2. 漏洞列表及修复措施', level=1)
        vulns = self._get_vulnerability_list(data)
        vuln_table = doc.add_table(rows=1 + len(vulns), cols=5)
        vuln_table.style = 'Table Grid'
        headers = ['漏洞类型', '端点/位置', '严重级别', '修复状态', '修复方式']
        for i, h in enumerate(headers):
            set_cell_text(vuln_table.cell(0, i), h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
            set_cell_shading(vuln_table.cell(0, i), "1a237e")

        sev_colors = {
            'critical': RGBColor(0xd3, 0x2f, 0x2f),
            'high': RGBColor(0xf5, 0x7c, 0x00),
            'medium': RGBColor(0xfb, 0xc0, 0x2d),
            'info': RGBColor(0x38, 0x8e, 0x3c),
        }
        for ri, v in enumerate(vulns):
            row_idx = ri + 1
            sev = v.get('severity', 'medium')
            fixed = v.get('is_fixed', 0)
            set_cell_text(vuln_table.cell(row_idx, 0), v.get('vuln_type', ''), size=10)
            set_cell_text(vuln_table.cell(row_idx, 1), v.get('endpoint', ''), size=10)
            set_cell_text(vuln_table.cell(row_idx, 2), sev.upper(), size=10,
                          color=sev_colors.get(sev, RGBColor(0x33, 0x33, 0x33)))
            set_cell_text(vuln_table.cell(row_idx, 3), '✓ 已修复' if fixed else '✗ 未修复', size=10,
                          color=RGBColor(0x38, 0x8e, 0x3c) if fixed else RGBColor(0xd3, 0x2f, 0x2f))
            set_cell_text(vuln_table.cell(row_idx, 4), v.get('fixed_method', '-'), size=10)

        doc.add_paragraph('')

        # ══════════════════════════════════
        # 3. 权限变更清单
        # ══════════════════════════════════
        doc.add_heading('3. 权限变更清单', level=1)

        perm_changes = data.get("permission_changes", [])
        tasks = data.get("tasks", [])
        completed_task_ids = {t.get("task_id") for t in tasks if t.get("status") == "completed"}
        perm_fixed = "fix_testuser_revoke" in completed_task_ids or "fix_devuser_revoke" in completed_task_ids or "fix_appuser_revoke" in completed_task_ids
        if not perm_changes:
            status_text = "已处理" if perm_fixed else "待处理"
            perm_changes = [
                {"change_type": "REVOKE", "target_user": "test_user", "target_host": "%", "privilege": "DELETE,DROP,FILE", "status": status_text},
                {"change_type": "DELETE", "target_user": "root(远程)", "target_host": "%", "privilege": "禁止远程登录", "status": status_text},
                {"change_type": "REVOKE", "target_user": "dev_user", "target_host": "%", "privilege": "DELETE,DROP,ALTER", "status": status_text},
                {"change_type": "REVOKE", "target_user": "app_user", "target_host": "192.168.%", "privilege": "DELETE", "status": status_text},
                {"change_type": "DROP", "target_user": "匿名用户", "target_host": "localhost", "privilege": "删除匿名用户", "status": status_text},
            ]

        if perm_changes:
            perm_table = doc.add_table(rows=1 + len(perm_changes), cols=5)
            perm_table.style = 'Table Grid'
            perm_headers = ['操作', '用户', 'Host', '权限/操作', '状态']
            for i, h in enumerate(perm_headers):
                set_cell_text(perm_table.cell(0, i), h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
                set_cell_shading(perm_table.cell(0, i), "1a237e")
            for ri, c in enumerate(perm_changes):
                row_idx = ri + 1
                set_cell_text(perm_table.cell(row_idx, 0), c.get('change_type', ''), size=10)
                set_cell_text(perm_table.cell(row_idx, 1), c.get('target_user', ''), size=10)
                set_cell_text(perm_table.cell(row_idx, 2), c.get('target_host', ''), size=10)
                set_cell_text(perm_table.cell(row_idx, 3), c.get('privilege', ''), size=10)
                st = c.get('status', '')
                sc = RGBColor(0x38, 0x8e, 0x3c) if st in ('applied', '已处理') else RGBColor(0xf5, 0x7c, 0x00)
                set_cell_text(perm_table.cell(row_idx, 4), st, size=10, color=sc)

        doc.add_paragraph('')

        # ══════════════════════════════════
        # 4. 数据恢复情况
        # ══════════════════════════════════
        doc.add_heading('4. 数据恢复情况', level=1)

        recovery_tasks = data.get("tasks", [])
        recovery_completed = {t.get("task_id") for t in recovery_tasks if t.get("status") == "completed"}
        recovery_done = all(tid in recovery_completed for tid in ["restore_full_backup", "apply_binlog_pitr", "verify_data_integrity"])
        recovery_partial = any(tid in recovery_completed for tid in ["restore_full_backup", "apply_binlog_pitr", "verify_data_integrity"])
        branch_name = data.get("student", {}).get("current_branch", "未知")

        p = doc.add_paragraph()
        run = p.add_run(f"恢复类型: ")
        run.bold = True
        p.add_run("时间点恢复 (PITR)")

        p = doc.add_paragraph()
        run = p.add_run("恢复工具: ")
        run.bold = True
        p.add_run("Percona XtraBackup + mysqlbinlog")

        if recovery_done:
            p = doc.add_paragraph()
            run = p.add_run("恢复状态: ")
            run.bold = True
            run2 = p.add_run("✓ 已完成")
            run2.font.color.rgb = RGBColor(0x38, 0x8e, 0x3c)
            doc.add_paragraph('恢复过程:', style='Heading 3')
            steps = [
                "✓ 使用 XtraBackup 准备全量备份 (2024-01-15_full.xb)",
                "✓ 恢复数据文件到 /var/lib/mysql/",
                "✓ 使用 mysqlbinlog 回放二进制日志到指定时间点",
                "✓ 校验数据完整性 (checksum验证通过)",
            ]
            for s in steps:
                doc.add_paragraph(s, style='List Bullet')

            doc.add_paragraph('恢复结果:', style='Heading 3')
            results = [
                "✓ 成功恢复 1500 条交易记录",
                "✓ 数据完整性校验通过",
                "⚠ 发现 3 条 account_no 字段异常数据",
            ]
            for r_item in results:
                p = doc.add_paragraph(r_item, style='List Bullet')
        elif recovery_partial:
            p = doc.add_paragraph()
            run = p.add_run("恢复状态: ")
            run.bold = True
            p.add_run("⚠ 部分完成")
            doc.add_paragraph('恢复过程:', style='Heading 3')
            recovery_steps_status = [
                ("全量备份恢复", "restore_full_backup" in recovery_completed),
                ("Binlog时间点恢复", "apply_binlog_pitr" in recovery_completed),
                ("数据完整性校验", "verify_data_integrity" in recovery_completed),
            ]
            for name, done in recovery_steps_status:
                icon = "✓" if done else "✗"
                p = doc.add_paragraph(f"{icon} {name}", style='List Bullet')
                if not done:
                    p.runs[0].font.color.rgb = RGBColor(0xd3, 0x2f, 0x2f)
            doc.add_paragraph("⚠ 数据恢复尚未全部完成，请继续执行剩余恢复步骤")
        else:
            p = doc.add_paragraph()
            run = p.add_run("恢复状态: ")
            run.bold = True
            p.add_run("未执行")
            doc.add_paragraph(f"学员当前处于「{branch_name}」分支，未选择数据恢复路线或未执行恢复操作。")
            doc.add_paragraph("如需数据恢复，请依次执行:", style='Heading 3')
            recovery_cmds = [
                "xtrabackup --prepare --apply-log-only --target-dir=/backup/full/",
                "xtrabackup --copy-back --target-dir=/backup/full/ --datadir=/var/lib/mysql/",
                "mysqlbinlog --stop-datetime='恢复时间点' /backup/binlog/mysql-bin.000012 | mysql -u root -p",
                "CHECKSUM TABLE core_bank.trade_flow 校验数据完整性",
            ]
            for cmd in recovery_cmds:
                doc.add_paragraph(cmd, style='List Bullet')

        doc.add_paragraph('')

        # ══════════════════════════════════
        # 5. 修复加固建议
        # ══════════════════════════════════
        doc.add_heading('5. 修复加固建议', level=1)

        recommendations = [
            ("1. 立即执行:", [
                "所有生产环境使用参数化查询，杜绝字符串拼接SQL",
                "撤销所有非管理员的全局DELETE/DROP/FILE权限",
                "禁止root远程登录，仅允许localhost访问",
                "删除匿名用户，启用密码强度验证插件",
            ]),
            ("2. 短期整改（1-2周）:", [
                "部署WAF (ModSecurity + OWASP CRS规则集)",
                "启用审计日志和慢查询日志",
                "修改默认MySQL端口(3306)",
                "启用SSL/TLS加密传输",
            ]),
            ("3. 长期规划:", [
                "部署数据库审计系统",
                "建立定期安全基线检查机制 (每月)",
                "实施数据库权限季度审查制度",
            ]),
        ]
        for title, items in recommendations:
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(12)
            for item in items:
                doc.add_paragraph(f"• {item}", style='List Bullet')

        doc.add_paragraph('')

        # ══════════════════════════════════
        # 6. 综合评分与等级
        # ══════════════════════════════════
        doc.add_heading('6. 综合评分与等级', level=1)

        score_rows_data = [
            ("基础分", f"{score_data['base_score']}/100"),
            ("失败扣分", f"-{score_data['fail_penalty']}"),
            ("重试扣分", f"-{score_data['retry_penalty']}"),
            ("漏洞修复加分", f"+{score_data['vuln_bonus']}"),
            ("最终得分", f"{score_data['final_score']}/100"),
            ("等级评定", score_data['grade']),
        ]
        score_table2 = doc.add_table(rows=len(score_rows_data), cols=2)
        score_table2.style = 'Table Grid'
        for i, (k, v) in enumerate(score_rows_data):
            is_final = k in ("最终得分", "等级评定")
            set_cell_text(score_table2.cell(i, 0), k, bold=is_final, size=11)
            set_cell_text(score_table2.cell(i, 1), v, bold=is_final, size=11)
            set_cell_shading(score_table2.cell(i, 0), "e8eaf6")

        doc.add_paragraph('')

        # ══════════════════════════════════
        # 7. 各分支错误次数
        # ══════════════════════════════════
        doc.add_heading('7. 各分支错误次数', level=1)
        branch_counts = self._get_branch_failure_counts(data)
        if branch_counts:
            rows = []
            for branch in ("baseline", "recovery", "sqli"):
                cnt = branch_counts.get(branch, 0)
                maxf = STORY_BRANCHES.get(branch, {}).get("max_failures", "-")
                title = STORY_BRANCHES.get(branch, {}).get("title", branch)
                rows.append((branch, title, f"{cnt} 次", f"{maxf} 次"))
            for branch, cnt in branch_counts.items():
                if branch not in ("baseline", "recovery", "sqli"):
                    rows.append((branch, branch, f"{cnt} 次", "-"))
            branch_table = doc.add_table(rows=1 + len(rows), cols=4)
            branch_table.style = 'Table Grid'
            for i, h in enumerate(['分支', '名称', '错误次数', '错误预算']):
                set_cell_text(branch_table.cell(0, i), h, bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
                set_cell_shading(branch_table.cell(0, i), "1a237e")
            for ri, row in enumerate(rows):
                row_idx = ri + 1
                for ci, val in enumerate(row):
                    set_cell_text(branch_table.cell(row_idx, ci), val, size=10)
        else:
            p = doc.add_paragraph("无任何分支产生操作失误")

        doc.add_paragraph('')
        doc.add_paragraph('')

        # ── 页脚 ──
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run(
            "本报告由 数据库安全运维智能体(Agent 1) 自动生成\n"
            "符合银行业安全运维标准"
        )
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 输出为 bytes
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    def _get_score_data(self, student_id: int) -> Dict[str, Any]:
        """获取评分数据"""
        from core.agent_orchestrator import get_orchestrator
        orchestrator = get_orchestrator(self._db)
        return orchestrator.calculate_final_score(student_id)

    # ──────────────────────────────────────────
    # 报告导出
    # ──────────────────────────────────────────

    def export_report(self, student_id: int = 1, format: str = "html", terminal_state: dict = None) -> Dict[str, Any]:
        """
        导出报告

        Args:
            student_id: 学员ID
            format: 输出格式 (html/text/docx)
            terminal_state: 终端权限状态（用于动态报告）

        Returns:
            报告内容和文件路径
        """
        # 确保输出目录存在
        output_dir = self._config.REPORT_OUTPUT_DIR
        os.makedirs(output_dir, exist_ok=True)

        ts = terminal_state or {}
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if format == "html":
            content = self.generate_html_report(student_id, ts)
            filename = f"数据库安全运维与加固报告_{timestamp}.html"
            filepath = os.path.join(output_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            preview = content[:200] + "..." if len(content) > 200 else content
        elif format == "docx":
            content_bytes = self.generate_docx_report(student_id, ts)
            filename = f"数据库安全运维与加固报告_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)
            with open(filepath, "wb") as f:
                f.write(content_bytes)
            preview = f"[Word文档] {len(content_bytes)} bytes"
        else:
            content = self.generate_text_report(student_id, ts)
            filename = f"数据库安全运维与加固报告_{timestamp}.txt"
            filepath = os.path.join(output_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            preview = content[:200] + "..." if len(content) > 200 else content

        return {
            "success": True,
            "format": format,
            "filename": filename,
            "filepath": filepath,
            "content_preview": preview,
        }

    def export_json(self, student_id: int = 1) -> Dict[str, Any]:
        """导出JSON格式数据"""
        data = self.collect_all_data(student_id)
        score_data = self._get_score_data(student_id)
        data["score"] = score_data

        output_dir = self._config.REPORT_OUTPUT_DIR
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = os.path.join(output_dir, f"report_data_{timestamp}.json")

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2, default=str)

        return {"success": True, "filepath": filepath}